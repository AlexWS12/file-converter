from pathlib import Path
from docx import Document


def docx_to_txt(src: Path, dst: Path):
    """Extract text from DOCX and write to a TXT file."""
    doc = Document(str(src))
    lines = [p.text for p in doc.paragraphs]
    text = "\n".join(lines).strip()
    dst.write_text(text, encoding="utf-8")


def txt_to_docx(src: Path, dst: Path):
    """Create a DOCX from a TXT file, each line a paragraph."""
    text = Path(src).read_text(encoding="utf-8")
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(dst))
